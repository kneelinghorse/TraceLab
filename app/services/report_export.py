"""Report export service that generates Markdown, PDF, or DOCX outputs."""
from __future__ import annotations

import io
from dataclasses import dataclass
from datetime import datetime, timezone
from html import escape as html_escape
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple

from jinja2 import Environment, FileSystemLoader

from app.models.mission_protocol import MissionProtocolDraft

try:  # Optional dependency used for DOCX generation
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover - exercised via error path
    DOCX_AVAILABLE = False

try:  # Optional dependency used for PDF generation
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

    REPORTLAB_AVAILABLE = True
except ImportError:  # pragma: no cover - exercised via error path
    REPORTLAB_AVAILABLE = False


class ReportExportError(RuntimeError):
    """Raised when report generation fails."""


class UnsupportedReportFormat(ReportExportError):
    """Raised when an unsupported export format is requested."""


@dataclass(frozen=True)
class ReportExportResult:
    """Represents a rendered report artifact."""

    filename: str
    media_type: str
    content: bytes


class ReportExportService:
    """Builds traceable reports for missions across multiple formats."""

    SUPPORTED_FORMATS: Tuple[str, ...] = ("md", "pdf", "docx")

    def __init__(self, *, template_dir: Path | None = None, template_name: str = "mission.md.j2") -> None:
        base_dir = template_dir or Path(__file__).resolve().parent.parent / "templates" / "reports"
        if not base_dir.exists():
            raise ReportExportError(f"Report template directory not found: {base_dir}")

        self.template_dir = base_dir
        self.template_name = template_name
        self.env = Environment(loader=FileSystemLoader(str(self.template_dir)), autoescape=False, trim_blocks=True, lstrip_blocks=True)
        self.env.filters["markdown_escape"] = self._markdown_escape
        self.template = self.env.get_template(self.template_name)

    def export(
        self,
        mission_payload: MissionProtocolDraft | Dict[str, Any],
        *,
        format: str = "md",
        completion_percentage: int | None = None,
    ) -> ReportExportResult:
        fmt = format.lower()
        if fmt not in self.SUPPORTED_FORMATS:
            raise UnsupportedReportFormat(f"Unsupported report format: {fmt}")

        context = self._build_context(mission_payload, completion_percentage)
        markdown = self.template.render(**context)
        slug = context["meta"]["slug"]

        if fmt == "md":
            return ReportExportResult(
                filename=f"{slug}.md",
                media_type="text/markdown; charset=utf-8",
                content=markdown.encode("utf-8"),
            )
        if fmt == "pdf":
            return ReportExportResult(
                filename=f"{slug}.pdf",
                media_type="application/pdf",
                content=self._markdown_to_pdf(markdown),
            )
        return ReportExportResult(
            filename=f"{slug}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            content=self._markdown_to_docx(markdown),
        )

    # ------------------------------------------------------------------
    # Context builders
    # ------------------------------------------------------------------
    def _build_context(
        self,
        mission_payload: MissionProtocolDraft | Dict[str, Any],
        completion_percentage: int | None,
    ) -> Dict[str, Any]:
        draft = mission_payload if isinstance(mission_payload, MissionProtocolDraft) else MissionProtocolDraft.model_validate(mission_payload)
        generated_at = datetime.now(tz=timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")
        completion = completion_percentage if completion_percentage is not None else 0
        meta = {
            "mission_id": draft.mission_id,
            "version": draft.version,
            "title": draft.title or draft.mission_id,
            "project_id": draft.project_id,
            "status_label": draft.status.replace("_", " ").title(),
            "owner": draft.owner,
            "completion_percentage": completion,
            "generated_at": generated_at,
            "tags": sorted(draft.tags) if draft.tags else [],
            "slug": self._slugify(draft.mission_id or f"mission-{generated_at}"),
        }

        research_statement = None
        if draft.research_statement:
            rs = draft.research_statement
            research_statement = {
                "topic": rs.topic,
                "objective": rs.objective,
                "scope": rs.scope,
                "audience": rs.audience or "—",
                "methodology": rs.methodology or "—",
                "success_metrics": ", ".join(rs.success_metrics) if rs.success_metrics else "",
                "risks": ", ".join(rs.risks) if rs.risks else "",
            }

        key_questions = [
            {
                "index": idx,
                "question": item.question,
                "status": item.status.replace("_", " ").title(),
                "answer": item.answer or "—",
                "confidence_display": self._format_confidence(item.confidence),
                "owner": item.owner or "—",
            }
            for idx, item in enumerate(draft.key_questions, start=1)
        ]

        evidence_rows = [
            {
                "index": idx,
                "label": f"EV-{idx:03d}",
                "evidence_id": item.evidence_id,
                "summary": item.summary,
                "summary_long": item.summary,
                "source": item.source,
                "source_type": item.source_type or "",
                "citation": self._format_citation(item.source, item.chunk_id, item.relevance_score),
                "chunk_id": item.chunk_id,
                "relevance_display": self._format_relevance(item.relevance_score),
            }
            for idx, item in enumerate(draft.evidence, start=1)
        ]

        synthesis_sections = self._build_synthesis_sections(draft)
        quality_checkpoints = self._build_quality_rows(draft)
        quality_summary = self._summarize_quality(quality_checkpoints)
        methodology = self._build_methodology(draft)

        return {
            "meta": meta,
            "research_statement": research_statement,
            "key_questions": key_questions,
            "evidence_rows": evidence_rows,
            "appendix_evidence": evidence_rows,
            "synthesis": {"sections": synthesis_sections},
            "quality_checkpoints": quality_checkpoints,
            "quality_summary": quality_summary,
            "methodology": methodology,
            "discussion_guide": draft.discussion_guide or [],
        }

    def _build_synthesis_sections(self, draft: MissionProtocolDraft) -> List[Dict[str, Any]]:
        sections: List[Tuple[str, Iterable[str]]] = []
        if draft.synthesis:
            synthesis = draft.synthesis
            sections = [
                ("Key Insights", synthesis.key_insights),
                ("Surprising Findings", synthesis.surprising_findings),
                ("Contradictions", synthesis.contradictory_information),
                ("Resolutions", synthesis.contradiction_resolutions),
                ("Recommendations", synthesis.recommendations),
                ("Next Steps", synthesis.next_steps),
            ]
        return [
            {"title": title, "entries": [item for item in items if item]}
            for title, items in sections
        ]

    def _build_quality_rows(self, draft: MissionProtocolDraft) -> List[Dict[str, Any]]:
        rows = []
        for checkpoint in draft.quality_checkpoints:
            rows.append(
                {
                    "gate": checkpoint.gate.replace("_", " ").title(),
                    "status_label": checkpoint.status.title(),
                    "notes": checkpoint.notes or "",
                    "status": checkpoint.status,
                }
            )
        return rows

    def _summarize_quality(self, checkpoints: List[Dict[str, Any]]) -> Dict[str, int]:
        summary = {"passed": 0, "failed": 0, "pending": 0}
        for checkpoint in checkpoints:
            status = checkpoint.get("status", "pending")
            if status == "pass":
                summary["passed"] += 1
            elif status == "fail":
                summary["failed"] += 1
            else:
                summary["pending"] += 1
        return summary

    def _build_methodology(self, draft: MissionProtocolDraft) -> Dict[str, Any] | None:
        if not draft.methodology_details:
            return None
        details = draft.methodology_details
        segments = []
        for segment in details.participant_segments:
            if segment.percentage is not None:
                pct = segment.percentage if segment.percentage <= 1 else segment.percentage / 100
                descriptor = f"{segment.segment} ({pct:.0%})"
            elif segment.count is not None:
                descriptor = f"{segment.segment} ({segment.count})"
            else:
                descriptor = segment.segment
            segments.append(descriptor)
        return {
            "total_participants": details.total_participants if details.total_participants is not None else "n/a",
            "segments": segments,
            "recruitment_method": details.recruitment_method,
            "validation_steps": details.validation_steps_completed or [],
            "artifacts": details.artifacts_verified or [],
            "notes": details.notes or "",
        }

    # ------------------------------------------------------------------
    # Rendering helpers
    # ------------------------------------------------------------------
    @staticmethod
    def _slugify(value: str) -> str:
        sanitized = "-".join(value.strip().split()) if value else "report"
        return "".join(char if char.isalnum() or char == "-" else "-" for char in sanitized).lower()

    @staticmethod
    def _format_confidence(confidence: float | None) -> str:
        if confidence is None:
            return "—"
        percentage = max(0.0, min(confidence, 1.0)) * 100
        return f"{percentage:.0f}%"

    @staticmethod
    def _format_relevance(score: float | None) -> str:
        if score is None:
            return "—"
        return f"{score:.2f}"

    @staticmethod
    def _format_citation(source: str, chunk_id: str | None, score: float | None) -> str:
        parts = [source]
        if chunk_id:
            parts.append(f"chunk {chunk_id}")
        if score is not None:
            parts.append(f"relevance {score:.2f}")
        return " · ".join(parts)

    @staticmethod
    def _markdown_escape(value: Any) -> str:
        if value is None:
            return "—"
        text = str(value)
        if not text.strip():
            return "—"
        escaped = text.replace("|", "\\|")
        escaped = escaped.replace("\r\n", "\n").replace("\n", "<br />")
        return escaped

    # ------------------------------------------------------------------
    # Format converters
    # ------------------------------------------------------------------
    def _markdown_to_docx(self, markdown_text: str) -> bytes:
        if not DOCX_AVAILABLE:
            raise ReportExportError("DOCX export requires python-docx. Install python-docx to continue.")
        document = DocxDocument()
        for block in self._parse_blocks(markdown_text):
            if block["type"] == "heading":
                level = min(block.get("level", 1), 4)
                document.add_heading(self._normalize_block_text(block["text"]), level=level)
            elif block["type"] == "bullet":
                for item in block["items"]:
                    document.add_paragraph(self._normalize_block_text(item), style="List Bullet")
            else:
                document.add_paragraph(self._normalize_block_text(block["text"]))
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _markdown_to_pdf(self, markdown_text: str) -> bytes:
        if not REPORTLAB_AVAILABLE:
            raise ReportExportError("PDF export requires reportlab. Install reportlab to continue.")
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=LETTER, leftMargin=54, rightMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()
        heading_styles = {
            1: styles["Heading1"],
            2: styles["Heading2"],
            3: styles["Heading3"],
            4: styles["Heading4"],
        }
        body_style = styles["BodyText"]
        elements: List[Any] = []

        for block in self._parse_blocks(markdown_text):
            block_type = block["type"]
            if block_type == "heading":
                level = min(block.get("level", 1), 4)
                text = html_escape(self._normalize_block_text(block["text"]))
                elements.append(Paragraph(text, heading_styles[level]))
            elif block_type == "bullet":
                items = [
                    ListItem(Paragraph(html_escape(self._normalize_block_text(item)), body_style))
                    for item in block["items"]
                ]
                elements.append(ListFlowable(items, bulletType="bullet", leftIndent=18))
            else:
                text = html_escape(self._normalize_block_text(block["text"])).replace("\n", "<br/>")
                elements.append(Paragraph(text, body_style))
            elements.append(Spacer(1, 8))

        doc.build(elements)
        buffer.seek(0)
        return buffer.getvalue()

    def _parse_blocks(self, markdown_text: str) -> List[Dict[str, Any]]:
        blocks: List[Dict[str, Any]] = []
        current_bullets: List[str] = []
        for line in markdown_text.splitlines():
            stripped = line.strip()
            if not stripped:
                if current_bullets:
                    blocks.append({"type": "bullet", "items": current_bullets[:]})
                    current_bullets.clear()
                continue
            if stripped == "---":
                if current_bullets:
                    blocks.append({"type": "bullet", "items": current_bullets[:]})
                    current_bullets.clear()
                continue
            if stripped.startswith("#"):
                if current_bullets:
                    blocks.append({"type": "bullet", "items": current_bullets[:]})
                    current_bullets.clear()
                level = len(stripped) - len(stripped.lstrip("#"))
                text = stripped[level:].strip()
                blocks.append({"type": "heading", "text": text, "level": level})
                continue
            if stripped.startswith("- "):
                current_bullets.append(stripped[2:].strip())
                continue
            if current_bullets:
                blocks.append({"type": "bullet", "items": current_bullets[:]})
                current_bullets.clear()
            blocks.append({"type": "paragraph", "text": stripped})
        if current_bullets:
            blocks.append({"type": "bullet", "items": current_bullets[:]})
        return blocks

    @staticmethod
    def _normalize_block_text(text: str) -> str:
        return text.replace("<br />", "\n").replace("<br/>", "\n").strip()
