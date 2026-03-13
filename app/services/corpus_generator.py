"""
Synthetic UX Research Corpus Generator

Builds a multi-format corpus populated with synthetic, locale-aware PII to
support Presidio evaluation and redaction tuning workflows.
"""

from __future__ import annotations

import csv
import json
import random
import textwrap
from collections.abc import Iterable
from datetime import datetime
from io import StringIO
from pathlib import Path
from typing import Any

from docx import Document
from faker import Faker
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


class CorpusGenerator:
    """Generates synthetic UX research documents with annotated PII."""

    def __init__(
        self,
        output_dir: str = "data/corpus",
        seed: int = 42,
        locales: Iterable[str] | None = None,
    ) -> None:
        """
        Args:
            output_dir: Directory where generated documents are stored.
            seed: Random seed for reproducibility.
            locales: Iterable of Faker locales to sample from. Defaults to en_US/en_GB.
        """

        self.output_dir = Path(output_dir)
        self.annotations_dir = self.output_dir / "annotations"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        for subdirectory in (
            "transcripts",
            "surveys",
            "personas",
            "test_notes",
            "briefs",
        ):
            (self.output_dir / subdirectory).mkdir(exist_ok=True)
        self.annotations_dir.mkdir(exist_ok=True)

        if locales:
            locale_list = tuple(locales)
        else:
            locale_list = ("en_US", "en_GB")

        if not locale_list:
            raise ValueError(
                "At least one locale must be provided for corpus generation."
            )

        self.seed = seed
        random.seed(seed)

        self.fakers: dict[str, Faker] = {}
        for locale in locale_list:
            faker_instance = Faker(locale)
            faker_instance.seed_instance(seed)
            self.fakers[locale] = faker_instance

        self.available_locales = tuple(self.fakers.keys())

    # --------------------------------------------------------------------------
    # Helper utilities
    # --------------------------------------------------------------------------

    def _select_locale(self, preferred: str | None = None) -> str:
        """Return a valid locale, defaulting to a random choice."""
        if preferred and preferred in self.fakers:
            return preferred
        return random.choice(self.available_locales)

    @staticmethod
    def _find_all(text: str, value: str) -> Iterable[int]:
        """Yield all start indices of a substring within a string."""
        start = text.find(value)
        while start != -1:
            yield start
            start = text.find(value, start + 1)

    def _annotate_text(
        self,
        text: str,
        entities: dict[str, Any],
        base_offset: int = 0,
    ) -> list[dict[str, Any]]:
        """
        Create annotations for PII entities found in text.

        Returns list of annotation dicts with start, end, entity_type, value.
        """

        annotations: list[dict[str, Any]] = []
        seen: set[tuple[str, int, int]] = set()

        def add_entity(entity_type: str, raw_value: Any) -> None:
            if raw_value in (None, ""):
                return
            value = str(raw_value)
            for start in self._find_all(text, value):
                absolute_start = start + base_offset
                key = (entity_type, absolute_start, absolute_start + len(value))
                if key in seen:
                    continue
                annotations.append(
                    {
                        "start": absolute_start,
                        "end": absolute_start + len(value),
                        "entity_type": entity_type,
                        "value": value,
                    }
                )
                seen.add(key)

        add_entity("PERSON", entities.get("PERSON"))
        add_entity("EMAIL_ADDRESS", entities.get("EMAIL_ADDRESS"))
        add_entity("PHONE_NUMBER", entities.get("PHONE_NUMBER"))
        add_entity("ORGANIZATION", entities.get("ORGANIZATION"))
        add_entity("CREDIT_CARD", entities.get("CREDIT_CARD"))
        add_entity("US_SSN", entities.get("US_SSN"))
        add_entity("AGE", entities.get("AGE"))
        add_entity("IP_ADDRESS", entities.get("IP_ADDRESS"))
        add_entity("PARTICIPANT_ID", entities.get("PARTICIPANT_ID"))
        add_entity("PROJECT_ID", entities.get("PROJECT_ID"))

        location = entities.get("LOCATION")
        if isinstance(location, dict):
            for key in ("full_address", "city", "state", "zipcode"):
                add_entity("LOCATION", location.get(key))

        date_time = entities.get("DATE_TIME")
        if date_time:
            add_entity("DATE_TIME", str(date_time)[:10])

        annotations.sort(key=lambda item: item["start"])
        return annotations

    def _generate_pii_entities(
        self, locale: str | None = None
    ) -> tuple[str, dict[str, Any]]:
        """Generate a consistent set of PII entities for a document."""
        locale = self._select_locale(locale)
        fake = self.fakers[locale]

        gender = random.choice(["M", "F"])
        first_name = (
            fake.first_name_male() if gender == "M" else fake.first_name_female()
        )
        last_name = fake.last_name()
        full_name = f"{first_name} {last_name}"

        location = {
            "city": fake.city(),
            "state": fake.state_abbr() if locale == "en_US" else fake.county(),
            "zipcode": fake.zipcode() if locale == "en_US" else fake.postcode(),
            "full_address": fake.address(),
        }

        participant_id = fake.bothify(text="PID-####-####").upper()
        project_id = fake.bothify(text="PROJ-????-####").upper()

        entities = {
            "PERSON": full_name,
            "FIRST_NAME": first_name,
            "LAST_NAME": last_name,
            "EMAIL_ADDRESS": f"{first_name.lower()}.{last_name.lower()}@{fake.domain_name()}",
            "PHONE_NUMBER": fake.phone_number(),
            "LOCATION": location,
            "DATE_TIME": fake.date_time_between(
                start_date="-2y", end_date="now"
            ).isoformat(),
            "AGE": random.randint(25, 65),
            "CREDIT_CARD": fake.credit_card_number(),
            "US_SSN": fake.ssn() if locale == "en_US" else None,
            "OCCUPATION": fake.job(),
            "ORGANIZATION": fake.company(),
            "PARTICIPANT_ID": participant_id,
            "PROJECT_ID": project_id,
        }

        return locale, entities

    def _register_document(
        self,
        corpus_metadata: dict[str, Any],
        file_path: Path,
        doc_type: str,
        locale: str,
        annotation_count: int,
    ) -> None:
        """Track document statistics in the corpus metadata."""
        corpus_metadata["files"].append(str(file_path))
        corpus_metadata["documents"].append(
            {
                "path": str(file_path),
                "doc_type": doc_type,
                "locale": locale,
                "annotation_count": annotation_count,
            }
        )
        corpus_metadata["document_counts"][doc_type] = (
            corpus_metadata["document_counts"].get(doc_type, 0) + 1
        )
        corpus_metadata["locale_counts"][locale] = (
            corpus_metadata["locale_counts"].get(locale, 0) + 1
        )
        corpus_metadata["total_annotations"] += annotation_count

    def _save_annotations(
        self,
        file_path: Path,
        annotations: list[dict[str, Any]],
        source_text: str,
        metadata: dict[str, Any] | None = None,
    ) -> None:
        """Persist annotations and source text for later evaluation."""
        file_path = Path(file_path)
        try:
            relative_path = file_path.relative_to(self.output_dir)
        except ValueError:
            relative_path = file_path.name

        annotation_payload: dict[str, Any] = {
            "document_path": str(relative_path),
            "document_absolute": str(file_path.resolve()),
            "generated_at": datetime.utcnow().isoformat() + "Z",
            "annotation_count": len(annotations),
            "annotations": annotations,
            "source_text": source_text,
        }

        if metadata:
            annotation_payload["metadata"] = metadata

        annotation_file = self.annotations_dir / f"{file_path.stem}_annotations.json"
        with open(annotation_file, "w", encoding="utf-8") as handle:
            json.dump(annotation_payload, handle, indent=2)

    # --------------------------------------------------------------------------
    # Document generators
    # --------------------------------------------------------------------------

    def generate_interview_transcript_txt(
        self,
        doc_id: str,
        locale: str | None = None,
    ) -> tuple[str, list[dict[str, Any]], str, dict[str, Any]]:
        """Generate a synthetic interview transcript in TXT format."""
        locale, entities = self._generate_pii_entities(locale)
        colleague_fake = self.fakers.get("en_US", self.fakers[locale])
        colleague_name = colleague_fake.name()

        transcript = f"""User Interview Transcript - Session {doc_id}
Interview Date: {entities["DATE_TIME"][:10]}
Participant ID: {entities["PARTICIPANT_ID"]}
Participant: {entities["PERSON"]}
Project Code: {entities["PROJECT_ID"]}
Email: {entities["EMAIL_ADDRESS"]}
Phone: {entities["PHONE_NUMBER"]}
Location: {entities["LOCATION"]["city"]}, {entities["LOCATION"]["state"]}

Interviewer: Thank you for taking the time to speak with us today. Can you start by telling us a bit about yourself?

Participant: Sure. My name is {entities["PERSON"]}, and I'm {entities["AGE"]} years old. I work as a {entities["OCCUPATION"]} at {entities["ORGANIZATION"]}, which is located in {entities["LOCATION"]["city"]}, {entities["LOCATION"]["state"]}. I'm currently participating in project {entities["PROJECT_ID"]}.

Interviewer: Can you tell us about your experience with the product?

Participant: I've been using it for about six months now. At first, I was a bit hesitant because my colleague {colleague_name} mentioned some issues, but overall it's been pretty good.

Interviewer: What would you say are the main pain points?

Participant: The main thing is the navigation. Sometimes I can't find what I'm looking for, especially when I'm working from {entities["LOCATION"]["city"]}. Also, I wish the mobile version had better support for my workflow.

Interviewer: Thank you for that feedback. Is there anything else you'd like to share?

Participant: Not really. You can reach me at {entities["EMAIL_ADDRESS"]} if you have any follow-up questions. My direct number is {entities["PHONE_NUMBER"]}.
"""

        annotations = self._annotate_text(transcript, entities)
        return transcript, annotations, locale, entities

    def generate_interview_transcript_docx(
        self,
        doc_id: str,
        locale: str | None = None,
    ) -> tuple[Path, list[dict[str, Any]], str, str, dict[str, Any]]:
        """Generate a synthetic interview transcript in DOCX format."""
        transcript, annotations, locale, entities = (
            self.generate_interview_transcript_txt(doc_id, locale)
        )

        document = Document()
        document.add_heading("User Interview Transcript", 0)
        for line in transcript.split("\n"):
            if line.strip():
                document.add_paragraph(line)

        output_path = self.output_dir / "transcripts" / f"transcript_{doc_id}.docx"
        document.save(str(output_path))

        return output_path, annotations, transcript, locale, entities

    def generate_survey_responses_csv(
        self,
        num_responses: int = 10,
        locale: str | None = None,
    ) -> tuple[Path, list[dict[str, Any]], str, list[str]]:
        """Generate synthetic survey responses in CSV format."""

        def serialize_row(row: list[Any]) -> str:
            buffer = StringIO()
            writer = csv.writer(buffer, lineterminator="\n")
            writer.writerow(row)
            return buffer.getvalue()

        header = [
            "response_id",
            "participant_id",
            "project_id",
            "timestamp",
            "age",
            "gender",
            "location",
            "occupation",
            "email",
            "phone",
            "rating",
            "feedback",
        ]

        locales_used: list[str] = []
        annotations: list[dict[str, Any]] = []
        lines: list[str] = [serialize_row(header)]
        offset = len(lines[0])

        for index in range(num_responses):
            row_locale, entities = self._generate_pii_entities(locale)
            locales_used.append(row_locale)
            gender = random.choice(["Male", "Female"])

            feedback = (
                f"I'm a {entities['AGE']}-year-old {gender.lower()} working as a {entities['OCCUPATION']} "
                f"in {entities['LOCATION']['city']}, {entities['LOCATION']['state']}. "
                f"The product has been useful for my work at {entities['ORGANIZATION']} under project {entities['PROJECT_ID']}. "
                f"Contact me at {entities['EMAIL_ADDRESS']} or {entities['PHONE_NUMBER']} for follow-up."
            )

            row = [
                f"R{index + 1:03d}",
                entities["PARTICIPANT_ID"],
                entities["PROJECT_ID"],
                entities["DATE_TIME"],
                entities["AGE"],
                gender,
                f"{entities['LOCATION']['city']}, {entities['LOCATION']['state']}",
                entities["OCCUPATION"],
                entities["EMAIL_ADDRESS"],
                entities["PHONE_NUMBER"],
                random.randint(1, 5),
                feedback,
            ]

            row_line = serialize_row(row)
            annotations.extend(
                self._annotate_text(row_line, entities, base_offset=offset)
            )
            offset += len(row_line)
            lines.append(row_line)

        csv_content = "".join(lines)
        output_path = self.output_dir / "surveys" / "survey_responses.csv"
        with open(output_path, "w", encoding="utf-8", newline="") as handle:
            handle.write(csv_content)

        return output_path, annotations, csv_content, locales_used

    def generate_user_persona_pdf(
        self,
        doc_id: str,
        locale: str | None = None,
    ) -> tuple[Path, list[dict[str, Any]], str, str, dict[str, Any]]:
        """Generate a synthetic user persona in PDF format."""
        locale, entities = self._generate_pii_entities(locale)

        persona_text = textwrap.dedent(
            f"""
            User Persona: {entities["PERSON"]}

            Background:
            {entities["PERSON"]} is a {entities["AGE"]}-year-old {entities["OCCUPATION"]} living in {entities["LOCATION"]["city"]}, {entities["LOCATION"]["state"]}.
            They work at {entities["ORGANIZATION"]} and have been in their current role for several years.

            Goals:
            - Improve productivity and workflow efficiency
            - Better integration with existing tools
            - Enhanced mobile experience

            Pain Points:
            - Difficulty navigating complex interfaces
            - Limited mobile support
            - Integration challenges with current systems

            Contact:
            Email: {entities["EMAIL_ADDRESS"]}
            Phone: {entities["PHONE_NUMBER"]}
            Location: {entities["LOCATION"]["full_address"]}
            Participant ID: {entities["PARTICIPANT_ID"]}
            Project Code: {entities["PROJECT_ID"]}
            """
        ).strip()

        output_path = self.output_dir / "personas" / f"persona_{doc_id}.pdf"
        styles = getSampleStyleSheet()
        story: list[Any] = []
        for line in persona_text.strip().split("\n"):
            if line.strip():
                story.append(Paragraph(line, styles["Normal"]))
                story.append(Spacer(1, 12))

        document = SimpleDocTemplate(str(output_path), pagesize=letter)
        document.build(story)

        annotations = self._annotate_text(persona_text, entities)
        return output_path, annotations, persona_text, locale, entities

    def generate_user_persona_docx(
        self,
        doc_id: str,
        locale: str | None = None,
    ) -> tuple[Path, list[dict[str, Any]], str, str, dict[str, Any]]:
        """Generate a synthetic user persona in DOCX format."""
        locale, entities = self._generate_pii_entities(locale)

        persona_content = textwrap.dedent(
            f"""
            Background:
            {entities["PERSON"]} is a {entities["AGE"]}-year-old {entities["OCCUPATION"]} living in {entities["LOCATION"]["city"]}, {entities["LOCATION"]["state"]}.
            They work at {entities["ORGANIZATION"]} and have been in their current role for several years.

            Goals:
            - Improve productivity and workflow efficiency
            - Better integration with existing tools

            Contact:
            Email: {entities["EMAIL_ADDRESS"]}
            Phone: {entities["PHONE_NUMBER"]}
            Location: {entities["LOCATION"]["full_address"]}
            Participant ID: {entities["PARTICIPANT_ID"]}
            Project Code: {entities["PROJECT_ID"]}
            """
        ).strip()

        document = Document()
        document.add_heading(f"User Persona: {entities['PERSON']}", 0)
        for line in persona_content.strip().split("\n"):
            if line.strip():
                document.add_paragraph(line)

        output_path = self.output_dir / "personas" / f"persona_{doc_id}.docx"
        document.save(str(output_path))

        annotations = self._annotate_text(persona_content, entities)
        return output_path, annotations, persona_content, locale, entities

    def generate_usability_test_notes_txt(
        self,
        doc_id: str,
        locale: str | None = None,
    ) -> tuple[str, list[dict[str, Any]], str, dict[str, Any]]:
        """Generate synthetic usability test notes in TXT format."""
        locale, entities = self._generate_pii_entities(locale)
        fake = self.fakers[locale]
        ip_address = fake.ipv4()
        entities["IP_ADDRESS"] = ip_address

        notes = textwrap.dedent(
            f"""
            Usability Test Notes - Session {doc_id}
            Date: {entities["DATE_TIME"][:10]}
            Participant ID: {entities["PARTICIPANT_ID"]}
            Participant: {entities["PERSON"]}
            Project Code: {entities["PROJECT_ID"]}
            Email: {entities["EMAIL_ADDRESS"]}

            Session Overview:
            Participant {entities["PERSON"]} (age {entities["AGE"]}, {entities["OCCUPATION"]}) from {entities["LOCATION"]["city"]}, {entities["LOCATION"]["state"]}
            tested the prototype interface.

            Key Observations:
            - Participant struggled with navigation menu
            - Expressed confusion about checkout flow
            - Positive feedback on visual design
            - Mentioned friend {fake.name()} had similar issues

            Technical Details:
            - Session IP: {ip_address}
            - Browser: Chrome
            - Device: Desktop
            - Recording URL: https://test.example.com/recordings/{doc_id}

            Follow-up:
            Contact {entities["PERSON"]} at {entities["EMAIL_ADDRESS"]} or {entities["PHONE_NUMBER"]} for debrief.
            """
        ).strip()

        annotations = self._annotate_text(notes, entities)
        return notes, annotations, locale, entities

    def generate_research_brief_markdown(
        self,
        doc_id: str,
        locale: str | None = None,
    ) -> tuple[str, list[dict[str, Any]], str, dict[str, Any]]:
        """Generate a synthetic research brief in Markdown format."""
        locale, entities = self._generate_pii_entities(locale)

        markdown = textwrap.dedent(
            f"""
            ---
            title: "UX Research Brief {doc_id}"
            author: "{entities["PERSON"]}"
            author_email: "{entities["EMAIL_ADDRESS"]}"
            interview_date: "{entities["DATE_TIME"][:10]}"
            locale: "{locale}"
            participant_id: "{entities["PARTICIPANT_ID"]}"
            project_id: "{entities["PROJECT_ID"]}"
            ---

            # Executive Summary
            Stakeholder {entities["PERSON"]} ({entities["OCCUPATION"]} at {entities["ORGANIZATION"]}) participated in a moderated study conducted in {entities["LOCATION"]["city"]}, {entities["LOCATION"]["state"]}. Key findings highlight workflow gaps that affect teams located near {entities["LOCATION"]["full_address"]}.

            # Participant Snapshot
            - Name: {entities["PERSON"]}
            - Participant ID: {entities["PARTICIPANT_ID"]}
            - Project Code: {entities["PROJECT_ID"]}
            - Email: {entities["EMAIL_ADDRESS"]}
            - Phone: {entities["PHONE_NUMBER"]}
            - Age: {entities["AGE"]}
            - Location: {entities["LOCATION"]["city"]}, {entities["LOCATION"]["state"]}

            # Key Quotes
            > "When I log in from {entities["LOCATION"]["city"]}, the dashboard buries the most critical metrics."

            > "Feel free to reach me at {entities["EMAIL_ADDRESS"]} if you need clarification."

            # Recommendations
            1. Streamline navigation for regional dashboards.
            2. Improve sign-in audit logging for IP addresses similar to the ones observed during testing.
            3. Provide guided tours targeted at organizations like {entities["ORGANIZATION"]}.
            """
        ).strip()

        annotations = self._annotate_text(markdown, entities)
        return markdown, annotations, locale, entities

    # --------------------------------------------------------------------------
    # Corpus generation orchestrator
    # --------------------------------------------------------------------------

    def generate_corpus(
        self,
        transcript_txt_count: int = 200,
        transcript_docx_count: int = 200,
        survey_responses: int = 300,
        persona_pdf_count: int = 75,
        persona_docx_count: int = 75,
        test_notes_count: int = 150,
        research_brief_count: int = 100,
    ) -> dict[str, Any]:
        """
        Generate the complete corpus.

        Returns corpus metadata including totals and per-document stats.
        """

        corpus_metadata: dict[str, Any] = {
            "generation_timestamp": datetime.utcnow().isoformat() + "Z",
            "seed": self.seed,
            "files": [],
            "documents": [],
            "document_counts": {},
            "locale_counts": {},
            "total_documents": 0,
            "total_annotations": 0,
            "parameters": {
                "transcript_txt_count": transcript_txt_count,
                "transcript_docx_count": transcript_docx_count,
                "survey_responses": survey_responses,
                "persona_pdf_count": persona_pdf_count,
                "persona_docx_count": persona_docx_count,
                "test_notes_count": test_notes_count,
                "research_brief_count": research_brief_count,
            },
        }

        doc_counter = 1

        # Interview transcripts (TXT)
        print(f"Generating {transcript_txt_count} interview transcripts (TXT)...")
        for _ in range(transcript_txt_count):
            doc_id = f"TXT{doc_counter:04d}"
            transcript, annotations, locale, entities = (
                self.generate_interview_transcript_txt(doc_id)
            )
            output_path = self.output_dir / "transcripts" / f"transcript_{doc_id}.txt"
            with open(output_path, "w", encoding="utf-8") as handle:
                handle.write(transcript)

            annotation_metadata = {
                "doc_type": "transcript_txt",
                "locale": locale,
                "entities": entities,
            }
            self._save_annotations(
                output_path, annotations, transcript, annotation_metadata
            )
            self._register_document(
                corpus_metadata, output_path, "transcript_txt", locale, len(annotations)
            )
            doc_counter += 1

        # Interview transcripts (DOCX)
        print(f"Generating {transcript_docx_count} interview transcripts (DOCX)...")
        for _ in range(transcript_docx_count):
            doc_id = f"DOCX{doc_counter:04d}"
            output_path, annotations, transcript, locale, entities = (
                self.generate_interview_transcript_docx(doc_id)
            )
            annotation_metadata = {
                "doc_type": "transcript_docx",
                "locale": locale,
                "entities": entities,
            }
            self._save_annotations(
                output_path, annotations, transcript, annotation_metadata
            )
            self._register_document(
                corpus_metadata,
                output_path,
                "transcript_docx",
                locale,
                len(annotations),
            )
            doc_counter += 1

        # Survey responses (CSV)
        print("Generating survey responses (CSV)...")
        output_path, annotations, csv_content, locales_used = (
            self.generate_survey_responses_csv(num_responses=survey_responses)
        )
        annotation_metadata = {
            "doc_type": "survey_csv",
            "locale": "mixed",
            "locales": locales_used,
            "records": survey_responses,
        }
        self._save_annotations(
            output_path, annotations, csv_content, annotation_metadata
        )
        self._register_document(
            corpus_metadata, output_path, "survey_csv", "mixed", len(annotations)
        )

        # User personas (PDF)
        print(f"Generating {persona_pdf_count} user personas (PDF)...")
        for _ in range(persona_pdf_count):
            doc_id = f"PDF{doc_counter:04d}"
            output_path, annotations, persona_text, locale, entities = (
                self.generate_user_persona_pdf(doc_id)
            )
            annotation_metadata = {
                "doc_type": "persona_pdf",
                "locale": locale,
                "entities": entities,
            }
            self._save_annotations(
                output_path, annotations, persona_text, annotation_metadata
            )
            self._register_document(
                corpus_metadata, output_path, "persona_pdf", locale, len(annotations)
            )
            doc_counter += 1

        # User personas (DOCX)
        print(f"Generating {persona_docx_count} user personas (DOCX)...")
        for _ in range(persona_docx_count):
            doc_id = f"PDC{doc_counter:04d}"
            output_path, annotations, persona_content, locale, entities = (
                self.generate_user_persona_docx(doc_id)
            )
            annotation_metadata = {
                "doc_type": "persona_docx",
                "locale": locale,
                "entities": entities,
            }
            self._save_annotations(
                output_path, annotations, persona_content, annotation_metadata
            )
            self._register_document(
                corpus_metadata, output_path, "persona_docx", locale, len(annotations)
            )
            doc_counter += 1

        # Usability test notes (TXT)
        print(f"Generating {test_notes_count} usability test notes (TXT)...")
        for _ in range(test_notes_count):
            doc_id = f"TST{doc_counter:04d}"
            notes, annotations, locale, entities = (
                self.generate_usability_test_notes_txt(doc_id)
            )
            output_path = self.output_dir / "test_notes" / f"test_notes_{doc_id}.txt"
            with open(output_path, "w", encoding="utf-8") as handle:
                handle.write(notes)

            annotation_metadata = {
                "doc_type": "test_notes_txt",
                "locale": locale,
                "entities": entities,
            }
            self._save_annotations(output_path, annotations, notes, annotation_metadata)
            self._register_document(
                corpus_metadata, output_path, "test_notes_txt", locale, len(annotations)
            )
            doc_counter += 1

        # Research briefs (Markdown)
        print(f"Generating {research_brief_count} research briefs (Markdown)...")
        for _ in range(research_brief_count):
            doc_id = f"MD{doc_counter:04d}"
            markdown, annotations, locale, entities = (
                self.generate_research_brief_markdown(doc_id)
            )
            output_path = self.output_dir / "briefs" / f"brief_{doc_id}.md"
            with open(output_path, "w", encoding="utf-8") as handle:
                handle.write(markdown)

            annotation_metadata = {
                "doc_type": "research_brief_md",
                "locale": locale,
                "entities": entities,
            }
            self._save_annotations(
                output_path, annotations, markdown, annotation_metadata
            )
            self._register_document(
                corpus_metadata,
                output_path,
                "research_brief_md",
                locale,
                len(annotations),
            )
            doc_counter += 1

        corpus_metadata["total_documents"] = len(corpus_metadata["files"])

        metadata_path = self.output_dir / "corpus_metadata.json"
        with open(metadata_path, "w", encoding="utf-8") as handle:
            json.dump(corpus_metadata, handle, indent=2)

        print("\nCorpus generation complete!")
        print(f"Total documents: {corpus_metadata['total_documents']}")
        print(f"Total annotations: {corpus_metadata['total_annotations']}")
        print(f"Output directory: {self.output_dir}")

        return corpus_metadata


if __name__ == "__main__":
    generator = CorpusGenerator(output_dir="data/corpus", seed=42)
    generator.generate_corpus(
        transcript_txt_count=200,
        transcript_docx_count=200,
        survey_responses=300,
        persona_pdf_count=75,
        persona_docx_count=75,
        test_notes_count=150,
        research_brief_count=100,
    )
